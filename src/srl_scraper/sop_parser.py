"""SOP（標準作業手順書）パーサー

病院の SOP（Word / PDF）から検査項目ごとの測定法・試薬・装置情報を抽出する。

対応形式:
  - .docx (Word)  → python-docx
  - .pdf (PDF)    → pdfplumber + pymupdf フォールバック
  - .xlsx / .csv  → 既存の仕組みで JSON 化済みを想定

SOP の典型的なセクション構造:
  1. 目的
  2. 測定原理 / 測定法  ← 抽出対象
  3. 試薬              ← 抽出対象
  4. 機器 / 装置        ← 抽出対象
  5. 操作手順
  6. 精度管理
  ...
"""

import json
import logging
import re
from datetime import datetime, timezone
from pathlib import Path

logger = logging.getLogger(__name__)

# セクション検出用キーワード（優先度順）
SECTION_KEYWORDS = {
    "title": [
        r"検査項目",
        r"項目名",
        r"測定項目",
    ],
    "method": [
        r"測定原理",
        r"測定法",
        r"測定方法",
        r"検査方法",
        r"原理",
        r"分析法",
    ],
    "reagent": [
        r"試薬",
        r"使用試薬",
        r"試薬[・／]",
        r"キット",
    ],
    "instrument": [
        r"分析装置",
        r"測定装置",
        r"使用機器",
        r"機器",
        r"装置",
        r"分析機",
    ],
    "purpose": [
        r"目的",
        r"臨床的意義",
    ],
    "specimen": [
        r"検体",
        r"材料",
        r"採取",
    ],
}


def _is_section_header(text: str) -> str | None:
    """テキストがセクション見出しかどうか判定し、カテゴリを返す

    見出しの特徴:
      - 短い（30文字以下）
      - 番号付き（「1.」「第1章」「Ⅰ.」「(1)」等）またはキーワードのみ
      - キーワードを含む
    """
    text = text.strip()
    if not text or len(text) > 40:
        return None

    # 見出しらしい構造かチェック（番号付き or キーワードのみの短文）
    has_number = bool(re.match(
        r"^[\d０-９]+[\.\．\)）]\s*|^第[\d０-９]+|^[ⅠⅡⅢⅣⅤ]+[\.\．]|^\([\d０-９]+\)",
        text,
    ))
    is_short_keyword = len(text) <= 15

    if not has_number and not is_short_keyword:
        return None

    for category, keywords in SECTION_KEYWORDS.items():
        for kw in keywords:
            if re.search(kw, text):
                return category
    return None


def _clean_section_text(lines: list[str]) -> str:
    """セクション内のテキストを整形"""
    cleaned = []
    for line in lines:
        line = line.strip()
        if line:
            cleaned.append(line)
    return "\n".join(cleaned)


# ---------------------------------------------------------------------------
# Word (.docx) パーサー
# ---------------------------------------------------------------------------

def parse_docx(filepath: Path) -> dict:
    """Word ファイルから SOP 情報を抽出"""
    import docx

    doc = docx.Document(str(filepath))

    # 全段落を取得（見出しスタイル情報付き）
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = para.style.name.startswith("Heading") if para.style else False
        paragraphs.append({
            "text": text,
            "style": para.style.name if para.style else "",
            "is_heading": is_heading,
        })

    # テーブルからもテキスト抽出
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                table_texts.append(" | ".join(row_texts))

    # Heading 1 をドキュメントタイトル（検査項目名）として抽出
    doc_title = ""
    for p in paragraphs:
        if p["style"] == "Heading 1":
            doc_title = p["text"]
            break

    # セクション分割
    sections = _split_into_sections(
        [p["text"] for p in paragraphs],
        [p["is_heading"] or len(p["text"]) < 50 for p in paragraphs],
    )

    # ドキュメントタイトルを検査項目名として補完
    if doc_title and not sections.get("title"):
        sections["title"] = doc_title

    # テーブル内容からも試薬・装置情報を探す
    for ttext in table_texts:
        cat = _is_section_header(ttext)
        if cat and cat not in sections:
            sections[cat] = ttext

    return {
        "source_file": str(filepath),
        "format": "docx",
        "sections": sections,
        "raw_paragraphs": [p["text"] for p in paragraphs],
        "tables": table_texts,
    }


# ---------------------------------------------------------------------------
# PDF パーサー
# ---------------------------------------------------------------------------

def parse_pdf(filepath: Path) -> dict:
    """PDF ファイルから SOP 情報を抽出"""
    texts = []

    # まず pdfplumber で試行（テーブル抽出に強い）
    try:
        import pdfplumber
        with pdfplumber.open(str(filepath)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    texts.extend(text.split("\n"))

                # テーブルも抽出
                for table in page.extract_tables():
                    for row in table:
                        row_text = " | ".join(cell or "" for cell in row if cell)
                        if row_text.strip():
                            texts.append(row_text)
    except Exception as e:
        logger.warning("pdfplumber 失敗: %s - %s", filepath, e)

    # テキストが取れなかった場合 pymupdf でフォールバック
    if not texts:
        try:
            import fitz
            doc = fitz.open(str(filepath))
            for page in doc:
                text = page.get_text()
                if text:
                    texts.extend(text.split("\n"))
            doc.close()
        except Exception as e:
            logger.warning("pymupdf 失敗: %s - %s", filepath, e)

    if not texts:
        logger.error("PDF からテキスト抽出できませんでした（画像 PDF の可能性）: %s", filepath)
        return {
            "source_file": str(filepath),
            "format": "pdf",
            "sections": {},
            "error": "テキスト抽出不可（画像PDF）",
            "raw_lines": [],
        }

    # セクション分割
    sections = _split_into_sections(texts)

    return {
        "source_file": str(filepath),
        "format": "pdf",
        "sections": sections,
        "raw_lines": texts,
    }


# ---------------------------------------------------------------------------
# 共通セクション分割ロジック
# ---------------------------------------------------------------------------

def _split_into_sections(
    lines: list[str],
    heading_hints: list[bool] | None = None,
) -> dict[str, str]:
    """行リストをセクションに分割して辞書で返す

    Returns:
        {"method": "測定原理テキスト...", "reagent": "試薬テキスト...", ...}
    """
    if heading_hints is None:
        heading_hints = [True] * len(lines)

    sections: dict[str, list[str]] = {}
    current_category: str | None = None

    for i, line in enumerate(lines):
        text = line.strip()
        if not text:
            continue

        # 見出し判定（短い行 or 見出しスタイル）
        if heading_hints[i] or len(text) < 60:
            cat = _is_section_header(text)
            if cat:
                current_category = cat
                # 見出し行自体もセクションに含める（項目名など）
                if cat == "title":
                    if cat not in sections:
                        sections[cat] = []
                    sections[cat].append(text)
                continue

        # 現在のセクションにテキストを追加
        if current_category:
            if current_category not in sections:
                sections[current_category] = []
            sections[current_category].append(text)

    return {k: _clean_section_text(v) for k, v in sections.items()}


# ---------------------------------------------------------------------------
# SOP → JSON 変換
# ---------------------------------------------------------------------------

def _extract_item_from_filename(filepath: str) -> str:
    """ファイル名から検査項目名を抽出

    パターン例:
      検査標準作業手順書_SO-生化学-001_総蛋白（TP）_第2版.docx
      → 総蛋白（TP）
    """
    name = Path(filepath).stem
    # _区切りで項目名部分を探す（SO-XXX-NNN の次）
    m = re.search(r"SO-[^_]+[-_]\d+[_\s]+(.+?)(?:_第\d+版|$)", name)
    if m:
        return m.group(1).strip()
    # 先頭の「検査標準作業手順書_」を除去
    name = re.sub(r"^検査標準作業手順書[_\s]*", "", name)
    return name


def extract_sop_info(parsed: dict) -> dict:
    """パース結果から構造化された SOP 情報を抽出"""
    sections = parsed.get("sections", {})

    # 検査項目名: 文書内 > ファイル名
    test_item = sections.get("title", "")
    if not test_item:
        test_item = _extract_item_from_filename(parsed.get("source_file", ""))

    # 測定法: 最初の1〜2行が重要（詳細な原理説明は補足）
    method_raw = sections.get("method", "")
    method_lines = method_raw.split("\n")
    method_summary = method_lines[0] if method_lines else ""
    method_detail = "\n".join(method_lines[1:]) if len(method_lines) > 1 else ""

    return {
        "source_file": parsed.get("source_file", ""),
        "format": parsed.get("format", ""),
        "test_item": test_item,
        "method_summary": method_summary,
        "method_detail": method_detail,
        "reagent": sections.get("reagent", ""),
        "instrument": sections.get("instrument", ""),
        "purpose": sections.get("purpose", ""),
        "specimen": sections.get("specimen", ""),
        "error": parsed.get("error", ""),
    }


def parse_sop(filepath: Path) -> dict:
    """ファイル形式を自動判定して SOP をパース"""
    suffix = filepath.suffix.lower()

    if suffix == ".docx":
        parsed = parse_docx(filepath)
    elif suffix == ".pdf":
        parsed = parse_pdf(filepath)
    else:
        raise ValueError(f"未対応の形式: {suffix} （.docx / .pdf のみ対応）")

    return extract_sop_info(parsed)


def parse_sop_directory(
    directory: Path,
    output_dir: Path | None = None,
) -> dict:
    """ディレクトリ内の全 SOP ファイルを一括パース"""
    output_dir = output_dir or Path("data")
    output_dir.mkdir(parents=True, exist_ok=True)

    results = []
    errors = []

    for filepath in sorted(directory.rglob("*")):
        if filepath.suffix.lower() not in (".docx", ".pdf"):
            continue
        if filepath.name.startswith("~"):
            continue
        logger.info("SOP パース: %s", filepath.name)
        try:
            info = parse_sop(filepath)
            results.append(info)
            status = "OK" if info["method_summary"] else "測定法なし"
            logger.info("  → %s (項目: %s)", status, info["test_item"][:30] or "不明")
        except Exception as e:
            logger.error("  → エラー: %s", e)
            errors.append({"file": str(filepath), "error": str(e)})

    now = datetime.now(timezone.utc)
    output = {
        "metadata": {
            "parsed_at": now.isoformat(),
            "source_directory": str(directory),
            "total_files": len(results),
            "with_method": sum(1 for r in results if r["method_summary"]),
            "errors": len(errors),
        },
        "sop_data": results,
        "errors": errors,
    }

    filepath = output_dir / "sop_parsed.json"
    filepath.write_text(json.dumps(output, ensure_ascii=False, indent=2), encoding="utf-8")
    logger.info("SOP パース完了: %s (%d件, 測定法あり%d件)",
                filepath, len(results), sum(1 for r in results if r["method_summary"]))

    return output
